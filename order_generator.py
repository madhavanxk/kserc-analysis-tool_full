"""
order_generator.py
==================
Generates a draft KSERC truing-up order as a Word (.docx) document.

Architecture: Chapterwise generation
- Numbers and regulatory citations injected deterministically from heuristic outputs
- Staff-approved amounts take precedence over system-calculated allowables
- Staff justifications become the Commission's stated reasoning in the order
- AI (Gemini) writes only the narrative paragraphs — never the numbers
- Returns: bytes (docx) for direct download in Streamlit
"""

import io
import json
import re
import subprocess
import tempfile
import os
import time
from datetime import datetime
from typing import Callable, Dict, Optional


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _get_approved(item: dict, fallback_key: str = 'allowable_value') -> float:
    """Return staff-approved amount if set, else system allowable, else 0."""
    sa = item.get('staff_approved_amount')
    if sa is not None:
        try:
            return float(sa)
        except (TypeError, ValueError):
            pass
    av = item.get(fallback_key) or item.get('allowable_value') or item.get('allowable') or 0
    try:
        return float(av)
    except (TypeError, ValueError):
        return 0.0


def _get_justification(item: dict) -> str:
    """Return staff justification if any, else empty string."""
    return (item.get('staff_justification') or '').strip()


def _get_review_status(item: dict) -> str:
    return item.get('staff_review_status', 'Pending')


def _flag_word(flag: str) -> str:
    return {'GREEN': 'accepted', 'YELLOW': 'conditionally accepted',
            'RED': 'rejected/reduced', 'GREY': 'noted'}.get(flag, 'noted')


def _gemini_call(prompt: str, api_key: str, max_tokens: int = 800) -> str:
    """Call Gemini Flash REST API. Returns text or empty string on failure."""
    import urllib.request
    url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
           f"gemini-2.0-flash:generateContent?key={api_key}")
    payload = json.dumps({
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"maxOutputTokens": max_tokens, "temperature": 0.2}
    }).encode()
    req = urllib.request.Request(url, data=payload,
                                  headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read())
            return data['candidates'][0]['content']['parts'][0]['text'].strip()
    except Exception:
        return ""


def _para_prompt(
    line_item: str,
    sbu: str,
    claimed: float,
    approved: float,
    flag: str,
    regulation: str,
    justification: str,
    calc_note: str = ""
) -> str:
    """Build the Gemini prompt for a single line item finding paragraph."""
    action = _flag_word(flag)
    just_clause = (f" Staff noting: {justification}" if justification else "")
    calc_clause = (f" Calculation note: {calc_note}" if calc_note else "")

    return (
        f"Write a 2-3 sentence regulatory finding paragraph in the style of a KSERC "
        f"truing-up order for the following line item. Use formal legal language. "
        f"Do NOT invent numbers — use only those provided.\n\n"
        f"SBU: {sbu}\n"
        f"Line Item: {line_item}\n"
        f"KSEB Claimed: Rs.{claimed:.2f} Cr\n"
        f"Commission Approved: Rs.{approved:.2f} Cr\n"
        f"Regulatory Basis: {regulation}\n"
        f"Decision: {action}\n"
        f"{just_clause}{calc_clause}\n\n"
        f"Output only the paragraph text, no headings or bullet points."
    )


# ─────────────────────────────────────────────────────────────────────────────
# DOCX BUILDER (pure Python via python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _build_docx(chapters: list) -> bytes:
    """
    Build Word document from list of chapter dicts.
    Each chapter: {'heading': str, 'sections': [{'title': str, 'paragraphs': [str],
                                                   'table': [(label, claimed, approved)]}]}
    Uses python-docx (available on Streamlit Cloud).
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed. Add 'python-docx' to requirements.txt.")

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Title page ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KERALA STATE ELECTRICITY REGULATORY COMMISSION")
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("DRAFT ORDER — TRUING UP OF ARR AND TARIFF").bold = True

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        f"In the matter of KSEB Ltd. Truing-Up Petition for FY 2024-25\n"
        f"Generated: {datetime.now().strftime('%d %B %Y')}"
    )

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = disclaimer.add_run(
        "\n⚠ AI-GENERATED DRAFT FOR INTERNAL REVIEW ONLY — NOT FOR PUBLICATION"
    )
    d_run.bold = True
    d_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_page_break()

    # ── Chapters ──
    for ch in chapters:
        # Chapter heading
        h = doc.add_heading(ch['heading'], level=1)
        h.runs[0].font.size = Pt(13)

        for sec in ch.get('sections', []):
            if sec.get('title'):
                sh = doc.add_heading(sec['title'], level=2)
                sh.runs[0].font.size = Pt(11)

            for para_text in sec.get('paragraphs', []):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.space_after = Pt(6)

            # Summary table if provided
            rows = sec.get('table', [])
            if rows:
                tbl = doc.add_table(rows=1 + len(rows), cols=3)
                tbl.style = 'Table Grid'
                # Header
                hdr = tbl.rows[0].cells
                for i, txt in enumerate(['Line Item', 'Claimed (Rs. Cr)', 'Approved (Rs. Cr)']):
                    hdr[i].text = txt
                    hdr[i].paragraphs[0].runs[0].bold = True
                # Data
                for i, (label, claimed, approved) in enumerate(rows):
                    row_cells = tbl.rows[i + 1].cells
                    row_cells[0].text = label
                    row_cells[1].text = f"{claimed:.2f}"
                    row_cells[2].text = f"{approved:.2f}"
                doc.add_paragraph()

    # ── Save to bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

_SBUG_REGULATIONS = {
    'om_expenses':        'Regulation 45, Annexure-7, Tariff Regulations 2021',
    'roe':                'Regulation 47, Tariff Regulations 2021',
    'depreciation':       'Regulation 48, Tariff Regulations 2021',
    'ifc':                'Regulations 29, 32, 34, Tariff Regulations 2021',
    'master_trust':       'Regulation 30, Tariff Regulations 2021',
    'fuel_costs':         'Regulation 51, Tariff Regulations 2021',
    'nti':                'Regulation 52, Tariff Regulations 2021',
    'intangible_assets':  'Regulation 49, Tariff Regulations 2021',
    'other_expenses':     'Regulation 53, Tariff Regulations 2021',
    'exceptional_items':  'Prudence assessment; Order Para 6.101-6.106',
}

_DISPLAY = {
    'om_expenses': 'O&M Expenses',
    'roe': 'Return on Equity',
    'depreciation': 'Depreciation',
    'ifc': 'Interest & Finance Charges',
    'master_trust': 'Master Trust Obligations',
    'fuel_costs': 'Fuel Costs',
    'nti': 'Non-Tariff Income',
    'intangible_assets': 'Intangible Assets Amortisation',
    'other_expenses': 'Other Expenses',
    'exceptional_items': 'Exceptional Items',
}


def _build_sbu_chapter(
    sbu_label: str,
    line_items: dict,
    display_map: dict,
    regulation_map: dict,
    api_key: str,
    progress_callback: Optional[Callable] = None,
    progress_start: int = 0,
    progress_end: int = 33,
) -> dict:
    """Build a chapter dict for one SBU."""
    sections = []
    total = len(line_items)
    table_rows = []

    for idx, (key, item) in enumerate(line_items.items()):
        if not isinstance(item, dict):
            continue
        # Only skip items explicitly marked as skipped or error
        if item.get('status') in ('skipped', 'error'):
            continue
        # Skip if no useful data at all
        claimed_check = (item.get('claimed_value') or item.get('claimed') or
                         item.get('primary_heuristic', {}).get('claimed_value') or 0)
        if claimed_check == 0 and not item.get('allowable_value') and not item.get('allowable'):
            continue

        name = display_map.get(key, key.replace('_', ' ').title())
        claimed = float(item.get('claimed_value') or
                        item.get('primary_heuristic', {}).get('claimed_value') or 0)
        approved = _get_approved(item)
        flag = item.get('flag') or item.get('staff_override_flag') or 'GREY'
        regulation = regulation_map.get(key, 'Tariff Regulations 2021')
        justification = _get_justification(item)

        # Calc note from primary heuristic steps (first 2 lines only)
        steps = (item.get('calculation_steps') or
                 item.get('primary_heuristic', {}).get('calculation_steps') or [])
        calc_note = ' | '.join(str(s) for s in steps[:2] if s)

        # Generate AI narrative or use template
        if api_key:
            prompt = _para_prompt(name, sbu_label, claimed, approved,
                                   flag, regulation, justification, calc_note)
            narrative = _gemini_call(prompt, api_key)
            time.sleep(0.3)
            if not narrative:  # API call failed — use template
                narrative = None
        else:
            narrative = None

        if not narrative:
            action = _flag_word(flag)
            narrative = (
                f"The Commission has examined the claim of Rs.{claimed:.2f} Cr "
                f"under {name} for {sbu_label}. After detailed analysis per "
                f"{regulation}, the Commission {action} the claim and approves "
                f"Rs.{approved:.2f} Cr."
            )

        # Staff override note
        if _get_review_status(item) == 'Overridden' and justification:
            narrative += f"\n\nThe Commission notes: {justification}"

        sections.append({
            'title': name,
            'paragraphs': [narrative],
            'table': []
        })

        table_rows.append((name, claimed, approved))

        if progress_callback and total > 0:
            pct = progress_start + int((idx + 1) / total * (progress_end - progress_start))
            progress_callback(pct, f"Drafting {sbu_label} — {name}...")

    # Summary table as last section
    sections.append({
        'title': f'{sbu_label} — Summary of Approved ARR',
        'paragraphs': [],
        'table': table_rows
    })

    total_claimed  = sum(r[1] for r in table_rows)
    total_approved = sum(r[2] for r in table_rows)
    table_rows.append(('TOTAL', total_claimed, total_approved))

    return {
        'heading': f'Chapter — {sbu_label}',
        'sections': sections
    }


def _build_sbu_t_chapter(sbu_t_data: dict, api_key: str,
                          progress_callback=None) -> dict:
    raw = sbu_t_data.get('line_items', [])
    if isinstance(raw, dict):
        items = {k: v for k, v in raw.items()
                 if isinstance(v, dict) and v.get('status') not in ('skipped', 'error')}
    else:
        items = {}
        for i, v in enumerate(raw):
            if isinstance(v, dict):
                k = (v.get('name') or v.get('heuristic_name') or
                     v.get('canonical') or f'item_{i}')
                normalised = dict(v)
                if 'claimed' in v and 'claimed_value' not in v:
                    normalised['claimed_value'] = v['claimed']
                if 'allowable' in v and 'allowable_value' not in v:
                    normalised['allowable_value'] = v['allowable']
                items[k] = normalised

    return _build_sbu_chapter(
        'SBU-T (Transmission)',
        items,
        {k: v.get('name', k.replace('_', ' ').title()) for k, v in items.items()},
        {},
        api_key,
        progress_callback,
        progress_start=34,
        progress_end=55
    )


def _build_sbu_d_chapter(sbu_d_data: dict, api_key: str,
                          progress_callback=None) -> dict:
    raw = sbu_d_data.get('line_items', [])
    if isinstance(raw, dict):
        items = {k: v for k, v in raw.items()
                 if isinstance(v, dict) and v.get('status') not in ('skipped', 'error')}
    else:
        # Flat list format — key by name
        items = {}
        for i, v in enumerate(raw):
            if isinstance(v, dict):
                k = (v.get('name') or v.get('heuristic_name') or
                     v.get('canonical') or f'item_{i}')
                # Normalise field names: list items use 'claimed'/'allowable'
                normalised = dict(v)
                if 'claimed' in v and 'claimed_value' not in v:
                    normalised['claimed_value'] = v['claimed']
                if 'allowable' in v and 'allowable_value' not in v:
                    normalised['allowable_value'] = v['allowable']
                items[k] = normalised

    return _build_sbu_chapter(
        'SBU-D (Distribution)',
        items,
        {k: v.get('name', k.replace('_', ' ').title()) for k, v in items.items()},
        {},
        api_key,
        progress_callback,
        progress_start=56,
        progress_end=80
    )


def _build_consolidated_chapter(results: dict) -> dict:
    """Chapter 4 — Consolidated ARR and revenue gap."""
    g_items = results.get('line_items', {})
    t_data  = results.get('sbu_t', {})
    d_data  = results.get('sbu_d', {})

    def _sum_approved(items):
        total = 0.0
        for v in items.values():
            if isinstance(v, dict) and v.get('status') not in ('skipped', 'error'):
                total += _get_approved(v)
        return total

    def _sum_list(lst):
        total = 0.0
        for item in lst:
            if isinstance(item, dict):
                total += _get_approved(item)
        return total

    g_approved = _sum_approved(g_items)

    t_raw = t_data.get('line_items', [])
    if isinstance(t_raw, dict):
        t_approved = _sum_approved(t_raw)
    else:
        t_approved = _sum_list(t_raw)

    d_raw = d_data.get('line_items', [])
    if isinstance(d_raw, dict):
        d_approved = _sum_approved(d_raw)
    else:
        d_approved = _sum_list(d_raw)

    total_approved = g_approved + t_approved + d_approved

    meta = results.get('metadata', {})
    fy   = meta.get('fiscal_year', '2024-25')

    summary_para = (
        f"Having examined the truing-up petition filed by KSEB Ltd for FY {fy} "
        f"and after detailed analysis of each line item across all three business units, "
        f"the Commission determines the total approved Annual Revenue Requirement as follows: "
        f"SBU-G (Generation) Rs.{g_approved:.2f} Cr; "
        f"SBU-T (Transmission) Rs.{t_approved:.2f} Cr; "
        f"SBU-D (Distribution) Rs.{d_approved:.2f} Cr; "
        f"giving a total approved ARR of Rs.{total_approved:.2f} Cr for FY {fy}."
    )

    return {
        'heading': 'Chapter — Consolidated ARR and Commission Directions',
        'sections': [
            {
                'title': 'Consolidated Approved ARR',
                'paragraphs': [summary_para],
                'table': [
                    ('SBU-G Generation',     g_approved,    g_approved),
                    ('SBU-T Transmission',   t_approved,    t_approved),
                    ('SBU-D Distribution',   d_approved,    d_approved),
                    ('TOTAL',                total_approved, total_approved),
                ]
            },
            {
                'title': 'Directions to KSEB Ltd',
                'paragraphs': [
                    "KSEB Ltd is hereby directed to: (1) implement the approved ARR as determined "
                    "in this order; (2) submit audited accounts for FY 2024-25 to the Commission "
                    "within 30 days of finalisation; (3) furnish the actuarial report, "
                    "Board-approved funding proposal, and State Government approval for Master Trust "
                    "obligations within 60 days as directed in Para 6.82 of OP No. 49/2024; "
                    "(4) ensure T&D loss reduction targets are met in FY 2025-26.",
                    "\n[STAFF NOTE: This directions section requires review by the Member "
                    "before finalisation. Add petition-specific directions as appropriate.]"
                ],
                'table': []
            }
        ]
    }


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def generate_order(
    results: dict,
    api_key: str = "",
    progress_callback: Optional[Callable] = None
) -> bytes:
    """
    Generate draft KSERC order as Word document bytes.

    Args:
        results:           Full analysis results dict from integration_pipeline.
        api_key:           Gemini API key (optional; falls back to template text).
        progress_callback: Optional fn(pct: int, msg: str) for Streamlit progress.

    Returns:
        bytes: .docx file content ready for st.download_button.
    """
    def _prog(pct, msg):
        if progress_callback:
            progress_callback(pct, msg)

    _prog(5, "Building SBU-G chapter...")

    # ── Chapter 1: SBU-G ──
    g_items = results.get('line_items', {})
    ch_g = _build_sbu_chapter(
        'SBU-G (Generation)',
        g_items,
        _DISPLAY,
        _SBUG_REGULATIONS,
        api_key,
        progress_callback,
        progress_start=5,
        progress_end=33
    )

    _prog(34, "Building SBU-T chapter...")

    # ── Chapter 2: SBU-T ──
    ch_t = _build_sbu_t_chapter(
        results.get('sbu_t', {}),
        api_key,
        progress_callback
    )

    _prog(56, "Building SBU-D chapter...")

    # ── Chapter 3: SBU-D ──
    ch_d = _build_sbu_d_chapter(
        results.get('sbu_d', {}),
        api_key,
        progress_callback
    )

    _prog(82, "Building consolidated chapter...")

    # ── Chapter 4: Consolidated ──
    ch_consol = _build_consolidated_chapter(results)

    _prog(88, "Assembling Word document...")

    chapters = [ch_g, ch_t, ch_d, ch_consol]
    docx_bytes = _build_docx(chapters)

    _prog(98, "Finalising...")

    return docx_bytes
